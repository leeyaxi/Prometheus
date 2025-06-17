import os
import json
import glob
import logging
import hashlib
from pathlib import Path
from typing import List, Set
from PIL import Image
import io
import fitz
import docx

import pytesseract
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    TextLoader, PDFMinerLoader, CSVLoader,
    UnstructuredExcelLoader, UnstructuredPowerPointLoader,
    Docx2txtLoader
)
from src.utils.common import BaseModule
from src.utils.data import *

logger = logging.getLogger(__name__)

SUFFIX_TO_LOADER = {
    ".txt": TextLoader,
    ".md": TextLoader,
    ".csv": CSVLoader,
    ".xlsx": UnstructuredExcelLoader,
    ".pptx": UnstructuredPowerPointLoader,
    ".docx": Docx2txtLoader,
}

# Windows系统，如果Tesseract没加入环境变量，需指定exe路径：
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


def ocr_image_from_bytes(image_bytes: bytes) -> str:
    """用 Tesseract 对图片字节做OCR，返回识别文本"""
    image = Image.open(io.BytesIO(image_bytes))
    # lang可以根据需求调整，如中文简体"chi_sim"，英文"eng"，组合用"+"连接
    text = pytesseract.image_to_string(image, lang="chi_sim+eng")
    return text


class DocumentLoader(BaseModule):
    def __init__(self, conf: dict):
        super().__init__(conf, "library")
        self.folder_path = self.module_conf.get("path", "docs")
        self.record_file = self.module_conf.get("record_file", "ingested.json")
        self.chunk_size = self.module_conf.get("chunk_size", 1000)
        self.chunk_overlap = self.module_conf.get("chunk_overlap", 200)
        self.enable_dedup = self.module_conf.get("deduplicate", True)

    def _get_hash(self, text: str) -> str:
        """计算文本 hash 用于去重"""
        return hashlib.md5(text.strip().encode("utf-8")).hexdigest()

    def _deduplicate_chunks(self, docs: List[Document]) -> List[Document]:
        """对 chunks 内容去重"""
        seen_hashes: Set[str] = set()
        unique_docs = []
        for doc in docs:
            text_hash = self._get_hash(doc.page_content)
            if text_hash not in seen_hashes:
                seen_hashes.add(text_hash)
                unique_docs.append(doc)
        return unique_docs

    def _extract_images_from_pdf(self, path: Path) -> List[str]:
        """从PDF中提取图片OCR文本"""
        all_texts = []
        try:
            pdf = fitz.open(str(path))
            for page_num in range(len(pdf)):
                page = pdf.load_page(page_num)
                images = page.get_images(full=True)  # full=True 确保拿到所有图片

                page_text = ""
                for img in images:
                    xref = img[0]
                    pix = fitz.Pixmap(pdf, xref)
                    if pix.n < 5:  # RGB or grayscale
                        img_bytes = pix.tobytes("png")
                    else:  # CMYK or other，需要转换
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                        img_bytes = pix.tobytes("png")
                    pix = None

                    image = Image.open(io.BytesIO(img_bytes))
                    text = pytesseract.image_to_string(image, lang='eng')
                    text = text_postprocess(text)
                    page_text += text + "\n"

                all_texts.append(f"--- Page {page_num + 1} ---\n" + page_text)
        except Exception as e:
            logger.warning(f"OCR extraction from PDF failed {path.name}: {e}")
        return all_texts

    def _extract_images_from_docx(self, path: Path) -> List[str]:
        """从DOCX中提取图片OCR文本"""
        texts = []
        try:
            doc = docx.Document(str(path))
            for rel in doc.part._rels:
                rel = doc.part._rels[rel]
                if "image" in rel.target_ref:
                    img_bytes = rel.target_part.blob
                    text = ocr_image_from_bytes(img_bytes)
                    texts.append(clean_docx_numbers(text))
        except Exception as e:
            logger.warning(f"OCR extraction from DOCX failed {path.name}: {e}")
        return texts

    def _clean_text(self, text: str) -> str:
        """简单文本清洗，去除多余空行和空白"""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return "\n".join(lines)

    def _load_file(self, path: Path) -> List[Document]:
        suffix = path.suffix.lower()
        loader_cls = SUFFIX_TO_LOADER.get(suffix)

        if suffix == ".pdf":
            ocr_texts = self._extract_images_from_pdf(path)
            docs = [Document(
                page_content=ocr_texts,
                metadata={"source": str(path), "type": "ocr"}
            )]
            return docs
        if not loader_cls:
            raise ValueError(f"Unsupported file type: {suffix}")

        docs = loader_cls(str(path)).load()
        # 后处理
        for doc in docs:
            doc.page_content = clean_chinese_docx_text(doc.page_content)
        if suffix == ".docx":
            ocr_texts = self._extract_images_from_docx(path)
            docs.append(Document(
                page_content=ocr_texts,
                metadata={"source": str(path), "type": "ocr"}
            ))
        return docs

    def load(self) -> List[Document]:
        docs = []
        raw_docs = []
        ingested_files = set()

        if os.path.exists(self.record_file):
            with open(self.record_file, 'r') as f:
                ingested_files = set(json.load(f))

        all_files = glob.glob(self.folder_path + "/**/**.*", recursive=True)
        new_files = []

        for file in all_files:
            path = Path(file)
            filename = path.name
            if filename in ingested_files:
                continue

            try:
                loaded = self._load_file(path)
                raw_docs.extend(loaded)
                new_files.append(filename)
                logger.info(f"✅ Loaded file: {filename}")
            except Exception as e:
                logger.warning(f"❌ Failed to load {filename}: {e}")

        with open(self.record_file, 'w') as f:
            json.dump(list(ingested_files.union(new_files)), f)

        if not raw_docs:
            logger.info("No new documents found.")
            return []

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap
        )

        for doc in raw_docs:
            try:
                docs.extend(splitter.split_documents([doc]))
            except Exception as e:
                logger.warning(f"Chunking failed: {e}")

        if self.enable_dedup:
            docs = self._deduplicate_chunks(docs)

        logger.info(f"Loaded {len(docs)} unique chunks from {len(new_files)} new files.")
        return docs
